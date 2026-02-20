#!/usr/bin/env python3
"""
Create binary seed files (PDF, DOCX, XLSX, images) for testing.
Run this once to generate the binary test files.
"""

import os

SEED_DIR = os.path.join(os.path.dirname(__file__), "files")


def create_pdf():
    """Create a sample PDF using pypdf."""
    from pypdf import PdfWriter
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(72, 700, "Agentic Filesystem Research Paper")
    c.setFont("Helvetica", 12)

    text = [
        "Abstract: This paper presents the design and implementation of the Agentic",
        "Filesystem, a tenant-scoped file storage and semantic search system designed",
        "for AI agent workflows. The system combines traditional file storage with",
        "modern vector search capabilities to enable intelligent document retrieval.",
        "",
        "1. Introduction",
        "As AI agents become more sophisticated, they require access to large corpora",
        "of documents for context-aware decision making. The Agentic Filesystem",
        "addresses this need by providing a unified API for file management and",
        "semantic search, enabling agents to store, retrieve, and query documents",
        "using natural language.",
        "",
        "2. Architecture",
        "The system consists of four main components:",
        "  - File API: RESTful endpoints for CRUD operations",
        "  - Search API: Semantic and hybrid search capabilities",
        "  - Indexing Pipeline: Async text extraction and embedding",
        "  - Vector Store: Qdrant-based vector index with tenant isolation",
        "",
        "3. Indexing Pipeline",
        "When a file is uploaded, the system automatically extracts text content",
        "using Apache Tika for binary formats. The extracted text is then chunked",
        "into 512-token segments with 10% overlap. Each chunk is embedded using",
        "OpenAI's text-embedding-3-small model and stored in Qdrant.",
        "",
        "4. Search Capabilities",
        "The system supports three search modes:",
        "  - Dense search: Pure vector similarity using cosine distance",
        "  - Sparse search: BM25-based keyword matching",
        "  - Hybrid search: Reciprocal Rank Fusion of dense and sparse results",
        "",
        "5. Results",
        "In our evaluation, hybrid search achieved 15% better recall compared to",
        "dense-only search, while maintaining sub-100ms query latency for collections",
        "up to 1 million documents.",
        "",
        "6. Conclusion",
        "The Agentic Filesystem provides a practical solution for AI agent document",
        "management, combining the reliability of traditional file storage with the",
        "power of modern semantic search.",
    ]

    y = 670
    for line in text:
        c.drawString(72, y, line)
        y -= 16
        if y < 72:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = 740

    c.save()

    os.makedirs(os.path.join(SEED_DIR, "pdfs"), exist_ok=True)
    with open(os.path.join(SEED_DIR, "pdfs", "research-paper.pdf"), "wb") as f:
        f.write(buf.getvalue())
    print("Created: pdfs/research-paper.pdf")


def create_simple_pdf():
    """Create a PDF without reportlab using basic PDF structure."""
    content = """%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 520 >>
stream
BT
/F1 18 Tf
72 700 Td
(Agentic Filesystem Research Paper) Tj
/F1 12 Tf
0 -30 Td
(Abstract: This paper presents the design of the Agentic Filesystem,) Tj
0 -16 Td
(a tenant-scoped file storage and semantic search system for AI agents.) Tj
0 -16 Td
(The system combines file storage with vector search for document retrieval.) Tj
0 -30 Td
(1. The Indexing Pipeline processes files through text extraction,) Tj
0 -16 Td
(chunking at 512 tokens with 10 percent overlap, and embedding.) Tj
0 -16 Td
(2. Hybrid search uses Reciprocal Rank Fusion for best results.) Tj
0 -16 Td
(3. Tenant isolation ensures complete data separation.) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000838 00000 n
trailer
<< /Size 6 /Root 1 0 R >>
startxref
921
%%EOF"""

    os.makedirs(os.path.join(SEED_DIR, "pdfs"), exist_ok=True)
    with open(os.path.join(SEED_DIR, "pdfs", "research-paper.pdf"), "wb") as f:
        f.write(content.encode("latin-1"))
    print("Created: pdfs/research-paper.pdf")


def create_invoice_pdf():
    """Create a simple invoice PDF."""
    content = """%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 480 >>
stream
BT
/F1 20 Tf
72 720 Td
(INVOICE) Tj
/F1 12 Tf
0 -40 Td
(Invoice Number: INV-2025-0042) Tj
0 -20 Td
(Date: February 15, 2025) Tj
0 -20 Td
(Customer: Acme AI Corporation) Tj
0 -30 Td
(Item: Agentic Platform License - Enterprise Tier) Tj
0 -16 Td
(Quantity: 1) Tj
0 -16 Td
(Unit Price: $2,500.00) Tj
0 -16 Td
(Item: Vector Database Hosting - Monthly) Tj
0 -16 Td
(Quantity: 12) Tj
0 -16 Td
(Unit Price: $150.00) Tj
0 -30 Td
(Subtotal: $4,300.00) Tj
0 -16 Td
(Tax: $344.00) Tj
0 -16 Td
(Total Due: $4,644.00) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000798 00000 n
trailer
<< /Size 6 /Root 1 0 R >>
startxref
881
%%EOF"""

    os.makedirs(os.path.join(SEED_DIR, "pdfs"), exist_ok=True)
    with open(os.path.join(SEED_DIR, "pdfs", "invoice.pdf"), "wb") as f:
        f.write(content.encode("latin-1"))
    print("Created: pdfs/invoice.pdf")


def create_png_image():
    """Create a simple PNG image with text."""
    try:
        from PIL import Image, ImageDraw, ImageFont

        img = Image.new("RGB", (600, 400), color=(255, 255, 255))
        draw = ImageDraw.Draw(img)

        draw.text((50, 30), "Agentic Platform Architecture", fill=(0, 0, 0))
        draw.text((50, 70), "Components:", fill=(50, 50, 50))
        draw.text((70, 100), "- File API: CRUD operations", fill=(50, 50, 50))
        draw.text((70, 130), "- Search API: Semantic + Hybrid", fill=(50, 50, 50))
        draw.text((70, 160), "- Indexing Pipeline: Async processing", fill=(50, 50, 50))
        draw.text((70, 190), "- Vector Store: Qdrant", fill=(50, 50, 50))
        draw.text((70, 220), "- File Store: Local / S3 / MinIO", fill=(50, 50, 50))
        draw.text((50, 270), "Tenant Isolation via payload filters", fill=(0, 0, 128))
        draw.text((50, 300), "Hybrid search with RRF fusion", fill=(0, 0, 128))

        # Draw some boxes
        draw.rectangle([40, 20, 560, 50], outline=(0, 100, 200), width=2)
        draw.rectangle([40, 60, 560, 240], outline=(100, 100, 100), width=1)
        draw.rectangle([40, 260, 560, 330], outline=(0, 0, 128), width=1)

        os.makedirs(os.path.join(SEED_DIR, "images"), exist_ok=True)
        img.save(os.path.join(SEED_DIR, "images", "diagram.png"))
        print("Created: images/diagram.png")
    except ImportError:
        print("Skipping PNG creation (Pillow not available)")
        create_minimal_png()


def create_minimal_png():
    """Create a minimal valid PNG without Pillow."""
    import struct
    import zlib

    width, height = 100, 50

    def create_chunk(chunk_type, data):
        chunk = chunk_type + data
        return struct.pack(">I", len(data)) + chunk + struct.pack(">I", zlib.crc32(chunk) & 0xFFFFFFFF)

    # PNG signature
    signature = b"\x89PNG\r\n\x1a\n"

    # IHDR
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    ihdr = create_chunk(b"IHDR", ihdr_data)

    # IDAT - simple white image with some blue pixels
    raw_data = b""
    for y in range(height):
        raw_data += b"\x00"  # filter byte
        for x in range(width):
            if 10 <= y <= 40 and 10 <= x <= 90:
                raw_data += b"\x00\x64\xc8"  # blue
            else:
                raw_data += b"\xff\xff\xff"  # white

    compressed = zlib.compress(raw_data)
    idat = create_chunk(b"IDAT", compressed)

    # IEND
    iend = create_chunk(b"IEND", b"")

    os.makedirs(os.path.join(SEED_DIR, "images"), exist_ok=True)
    with open(os.path.join(SEED_DIR, "images", "diagram.png"), "wb") as f:
        f.write(signature + ihdr + idat + iend)
    print("Created: images/diagram.png (minimal)")


def create_jpg_image():
    """Create a simple JPEG image."""
    try:
        from PIL import Image, ImageDraw

        img = Image.new("RGB", (400, 300), color=(240, 240, 245))
        draw = ImageDraw.Draw(img)

        draw.text((30, 20), "Dashboard Screenshot", fill=(0, 0, 0))
        draw.text((30, 50), "Files: 42 indexed", fill=(50, 50, 50))
        draw.text((30, 70), "Search queries: 128 today", fill=(50, 50, 50))
        draw.text((30, 90), "Agents: 3 active", fill=(50, 50, 50))
        draw.text((30, 130), "Recent uploads:", fill=(0, 0, 0))
        draw.text((50, 155), "report.pdf - 2.3 MB - indexed", fill=(0, 128, 0))
        draw.text((50, 175), "data.xlsx - 450 KB - indexed", fill=(0, 128, 0))
        draw.text((50, 195), "notes.txt - 12 KB - indexed", fill=(0, 128, 0))

        os.makedirs(os.path.join(SEED_DIR, "images"), exist_ok=True)
        img.save(os.path.join(SEED_DIR, "images", "screenshot.jpg"), "JPEG", quality=85)
        print("Created: images/screenshot.jpg")
    except ImportError:
        print("Skipping JPEG creation (Pillow not available)")


def create_docx():
    """Create a sample DOCX file."""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Quarterly Business Report", level=1)
        doc.add_heading("Q4 2024 Performance Summary", level=2)

        doc.add_paragraph(
            "This report summarizes the key performance indicators and achievements "
            "for the Agentic Platform during Q4 2024. The platform saw significant "
            "growth in both user adoption and technical capabilities."
        )

        doc.add_heading("Key Metrics", level=2)
        doc.add_paragraph("Total Files Processed: 1.2 million")
        doc.add_paragraph("Average Query Latency: 85ms")
        doc.add_paragraph("Search Accuracy (MRR@10): 0.82")
        doc.add_paragraph("Active Tenants: 47")
        doc.add_paragraph("Uptime: 99.97%")

        doc.add_heading("Technical Achievements", level=2)
        doc.add_paragraph(
            "The hybrid search implementation combining dense vector similarity with "
            "BM25 sparse retrieval was deployed in October. This resulted in a 15% "
            "improvement in search recall compared to dense-only search. The Reciprocal "
            "Rank Fusion algorithm effectively merged results from both retrieval methods."
        )

        doc.add_heading("Revenue", level=2)
        doc.add_paragraph(
            "Q4 revenue reached $1.2M, representing a 35% quarter-over-quarter increase. "
            "Enterprise tier subscriptions grew by 40%, driven by the new RAG capabilities "
            "and multi-tenant isolation features."
        )

        doc.add_heading("Roadmap", level=2)
        doc.add_paragraph(
            "Q1 2025 priorities include: API key authentication and rate limiting, "
            "file versioning support, expanded binary format support (CAD, video), "
            "and a web-based document viewer."
        )

        os.makedirs(os.path.join(SEED_DIR, "office"), exist_ok=True)
        doc.save(os.path.join(SEED_DIR, "office", "report.docx"))
        print("Created: office/report.docx")
    except ImportError:
        print("Skipping DOCX creation (python-docx not available)")


def create_xlsx():
    """Create a sample XLSX file."""
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "Revenue Data"

        headers = ["Month", "Revenue", "Users", "Files Processed", "Queries"]
        ws.append(headers)

        data = [
            ["2024-10", 350000, 12000, 380000, 95000],
            ["2024-11", 400000, 14500, 420000, 110000],
            ["2024-12", 450000, 16000, 400000, 125000],
            ["2025-01", 520000, 18000, 460000, 140000],
        ]
        for row in data:
            ws.append(row)

        ws2 = wb.create_sheet("Agent Performance")
        ws2.append(["Agent ID", "Tasks Completed", "Avg Duration (s)", "Success Rate"])
        ws2.append(["agent-001", 1250, 4.2, 0.97])
        ws2.append(["agent-002", 980, 3.8, 0.95])
        ws2.append(["agent-003", 1100, 5.1, 0.93])

        os.makedirs(os.path.join(SEED_DIR, "office"), exist_ok=True)
        wb.save(os.path.join(SEED_DIR, "office", "data.xlsx"))
        print("Created: office/data.xlsx")
    except ImportError:
        print("Skipping XLSX creation (openpyxl not available)")


if __name__ == "__main__":
    print("Creating binary seed files...")
    create_simple_pdf()
    create_invoice_pdf()
    create_minimal_png()
    create_docx()
    create_xlsx()
    create_jpg_image()
    print("Done!")
